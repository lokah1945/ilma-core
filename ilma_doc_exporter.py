#!/usr/bin/env python3
"""
ILMA Document Exporter v1.0  (2026-06-01)
=========================================
Markdown (canonical) -> DOCX / PDF / HTML. Embeds headings, TOC, figures,
citations. Free, offline (python-docx, reportlab, markdown, Pillow).

API:
  export(markdown_text, out_basepath, formats=("docx","pdf","html"),
         title=None, figures=None) -> dict {fmt: path}
"""
from __future__ import annotations
import re, os, html, unicodedata
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

# ── Unicode font (PDF): ensures ALL characters render, never boxes ────────────
_UNICODE_FONT_CANDIDATES = [
    ("DejaVuSans", "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
                   "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf",
                   "/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf",
                   "/usr/share/fonts/truetype/dejavu/DejaVuSans-BoldOblique.ttf"),
    ("NotoSans", "/usr/share/fonts/truetype/noto/NotoSans-Regular.ttf",
                 "/usr/share/fonts/truetype/noto/NotoSans-Bold.ttf",
                 "/usr/share/fonts/truetype/noto/NotoSans-Italic.ttf",
                 "/usr/share/fonts/truetype/noto/NotoSans-BoldItalic.ttf"),
    ("LiberationSans", "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
                       "/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf",
                       "/usr/share/fonts/truetype/liberation/LiberationSans-Italic.ttf",
                       "/usr/share/fonts/truetype/liberation/LiberationSans-BoldItalic.ttf"),
]
_PDF_FONT = None
_PDF_FONT_BOLD = None
_PDF_FONT_ITALIC = None
_FONT_CMAP = None


def _register_unicode_font():
    global _PDF_FONT, _PDF_FONT_BOLD, _PDF_FONT_ITALIC, _FONT_CMAP
    if _PDF_FONT is not None:
        return _PDF_FONT
    try:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.pdfbase.pdfmetrics import registerFontFamily
        for name, reg, bold, ital, boldital in _UNICODE_FONT_CANDIDATES:
            if not os.path.exists(reg):
                continue
            pdfmetrics.registerFont(TTFont(name, reg))
            b, i, bi = name + "-Bold", name + "-Italic", name + "-BoldItalic"
            pdfmetrics.registerFont(TTFont(b, bold if os.path.exists(bold) else reg))
            pdfmetrics.registerFont(TTFont(i, ital if os.path.exists(ital) else reg))
            pdfmetrics.registerFont(TTFont(bi, boldital if os.path.exists(boldital) else reg))
            registerFontFamily(name, normal=name, bold=b, italic=i, boldItalic=bi)
            _PDF_FONT, _PDF_FONT_BOLD, _PDF_FONT_ITALIC = name, b, i
            try:
                from fontTools.ttLib import TTFont as _FT
                ft = _FT(reg); cmap = set()
                for table in ft["cmap"].tables:
                    cmap.update(table.cmap.keys())
                _FONT_CMAP = cmap; ft.close()
            except Exception:
                _FONT_CMAP = None
            return _PDF_FONT
    except Exception:
        pass
    return None


_SYMBOL_FALLBACK = {
    "\u2014": "-", "\u2013": "-", "\u2012": "-", "\u2015": "-",
    "\u2018": "'", "\u2019": "'", "\u201a": "'", "\u201b": "'",
    "\u201c": '"', "\u201d": '"', "\u201e": '"', "\u201f": '"',
    "\u2026": "...", "\u2022": "-", "\u00b7": "-", "\u25cf": "-", "\u25aa": "-",
    "\u00a0": " ", "\u2009": " ", "\u200a": " ", "\u202f": " ", "\u2007": " ",
    "\u2192": "->", "\u2190": "<-", "\u21d2": "=>", "\u2261": "=",
    "\ufeff": "", "\u200b": "", "\u200c": "", "\u200d": "",
    "\u2212": "-", "\u00d7": "x", "\u00f7": "/",
}


def _sanitize_text(text: str, ensure_font: bool = True) -> str:
    if not text:
        return ""
    t = unicodedata.normalize("NFC", str(text))
    for k, v in _SYMBOL_FALLBACK.items():
        if k in t:
            t = t.replace(k, v)
    if ensure_font and _FONT_CMAP:
        out = []
        for ch in t:
            cp = ord(ch)
            if cp < 32 and ch not in "\n\t":
                continue
            if cp < 128 or cp in _FONT_CMAP:
                out.append(ch)
            else:
                fb = _SYMBOL_FALLBACK.get(ch)
                if fb is not None:
                    out.append(fb)
                else:
                    try:
                        out.append("[" + unicodedata.name(ch).title().replace(" ", "") + "]")
                    except ValueError:
                        out.append("")
        t = "".join(out)
    return t



def _is_table_sep(line: str) -> bool:
    """A GFM table separator row: | --- | :--: | ... |"""
    t = line.strip()
    if "|" not in t or "-" not in t:
        return False
    cells = [c.strip() for c in t.strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{2,}:?", c or "") for c in cells if c != "")


def _split_row(line: str) -> List[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def _parse_md(md: str) -> List[Tuple[str, Any]]:
    """Block parser -> (kind, payload). kinds: h1,h2,h3,li,p,img,table,blank."""
    blocks: List[Tuple[str, Any]] = []
    lines = md.splitlines()
    i = 0
    n = len(lines)
    while i < n:
        s = lines[i].rstrip()
        stripped = s.strip()
        # ── GFM table: header row, separator row, then body rows ──
        if "|" in stripped and i + 1 < n and _is_table_sep(lines[i + 1]):
            header = _split_row(s)
            rows = [header]
            j = i + 2
            while j < n and "|" in lines[j] and lines[j].strip():
                rows.append(_split_row(lines[j]))
                j += 1
            blocks.append(("table", rows))
            i = j
            continue
        if not stripped:
            blocks.append(("blank", ""))
        elif s.startswith("### "):
            blocks.append(("h3", s[4:]))
        elif s.startswith("## "):
            blocks.append(("h2", s[3:]))
        elif s.startswith("# "):
            blocks.append(("h1", s[2:]))
        elif re.match(r"^\s*[-*]\s+", s):
            blocks.append(("li", re.sub(r"^\s*[-*]\s+", "", s)))
        elif re.match(r"!\[.*\]\(.*\)", s):
            m = re.match(r"!\[(.*?)\]\((.*?)\)", s)
            blocks.append(("img", (m.group(2), m.group(1))))
        else:
            blocks.append(("p", s))
        i += 1
    return blocks


def _strip_inline(t: str) -> str:
    if isinstance(t, tuple):
        return t[0]
    # NFC-normalize + drop control chars so DOCX/HTML never get unreadable bytes
    t = unicodedata.normalize("NFC", str(t))
    t = "".join(ch for ch in t if ord(ch) >= 32 or ch in "\n\t")
    t = re.sub(r"\*\*(.*?)\*\*", r"\1", t)
    t = re.sub(r"\*(.*?)\*", r"\1", t)
    t = re.sub(r"`(.*?)`", r"\1", t)
    t = re.sub(r"\[(.*?)\]\((.*?)\)", r"\1", t)
    return t


def to_html(md: str, title: str = "Document", doc_type: str = "report") -> str:
    import markdown as mdlib
    body = mdlib.markdown(md, extensions=["tables", "fenced_code", "toc"])
    _fp = _fmt_profile(doc_type)
    _align = _fp.get("align", "justify")
    _indent = float(_fp.get("first_line_indent", 0) or 0)
    _block = bool(_fp.get("block_paragraphs"))
    _ls = _fp.get("line_spacing", 1.5)
    _pcss = (f"text-align:{_align};line-height:{_ls};" +
             ("margin:0 0 1em 0;" if _block else f"margin:0;text-indent:{_indent/72.0:.2f}in;"))
    return f"""<!DOCTYPE html><html><head><meta charset="utf-8">
<title>{html.escape(title)}</title>
<style>
body{{font-family:Georgia,serif;max-width:820px;margin:40px auto;line-height:1.6;color:#222;padding:0 20px}}
p{{{_pcss}}}
table{{border-collapse:collapse;width:100%;margin:1em 0}} th{{background:#1f4e79;color:#fff;padding:8px;text-align:left}} td{{border:1px solid #ccc;padding:8px}} tr:nth-child(even){{background:#f2f2f2}}
h1{{font-size:2em;border-bottom:2px solid #333;padding-bottom:.2em}}
h2{{font-size:1.5em;margin-top:1.5em;border-bottom:1px solid #ccc}}
h3{{font-size:1.2em}} img{{max-width:100%;height:auto}}
code{{background:#f4f4f4;padding:2px 5px;border-radius:3px}}
table{{border-collapse:collapse;width:100%}} td,th{{border:1px solid #ccc;padding:6px}}
.refs{{font-size:.9em}}
</style></head><body>
{body}
</body></html>"""



def _fmt_profile(doc_type):
    try:
        import sys as _sys
        _sys.path.insert(0, "/root/.hermes/profiles/ilma")
        from ilma_writing_templates import get_format_profile
        return get_format_profile(doc_type or "report")
    except Exception:
        return {"align": "justify", "first_line_indent": 0, "para_space_after": 8,
                "line_spacing": 1.15, "body_font": "serif", "body_size": 11,
                "page_size": "A4", "margin_in": 1.0, "block_paragraphs": True}


def _page_size(name):
    from reportlab.lib.pagesizes import A4, LETTER
    return LETTER if str(name).upper() == "LETTER" else A4


def to_docx(md: str, out_path: str, title: str = None,
            figures: Optional[Dict[str, str]] = None, doc_type: str = "report",
            doc_meta: Optional[Dict[str, str]] = None) -> str:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    fp = _fmt_profile(doc_type)
    _ALIGN = {"justify": WD_ALIGN_PARAGRAPH.JUSTIFY, "left": WD_ALIGN_PARAGRAPH.LEFT,
              "center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT}
    body_align = _ALIGN.get(fp.get("align", "justify"), WD_ALIGN_PARAGRAPH.JUSTIFY)
    indent_pt = float(fp.get("first_line_indent", 0) or 0)
    space_after = float(fp.get("para_space_after", 8) or 0)
    line_spacing = float(fp.get("line_spacing", 1.15) or 1.15)
    body_size = float(fp.get("body_size", 11) or 11)
    body_font = "Times New Roman" if fp.get("body_font") == "serif" else "Calibri"

    doc = Document()
    # page setup
    try:
        sec = doc.sections[0]
        mt = fp.get("margin_top_in", fp.get("margin_in", 1.0))
        ml = fp.get("margin_left_in", fp.get("margin_in", 1.0))
        mr = fp.get("margin_right_in", fp.get("margin_in", 1.0))
        mb = fp.get("margin_bottom_in", fp.get("margin_in", 1.0))
        sec.top_margin = Inches(mt); sec.bottom_margin = Inches(mb)
        sec.left_margin = Inches(ml); sec.right_margin = Inches(mr)
        if str(fp.get("page_size", "A4")).upper() == "A4":
            sec.page_width = Inches(8.27); sec.page_height = Inches(11.69)
    except Exception:
        pass
    # default style font
    try:
        st = doc.styles["Normal"]
        st.font.name = body_font; st.font.size = Pt(body_size)
        st.paragraph_format.line_spacing = line_spacing
    except Exception:
        pass

    def _style_body(par, is_first_of_block=True):
        pf = par.paragraph_format
        pf.alignment = body_align
        pf.line_spacing = line_spacing
        if fp.get("block_paragraphs"):
            pf.space_after = Pt(space_after); pf.first_line_indent = Inches(0)
        else:
            pf.space_after = Pt(0)
            pf.first_line_indent = Inches(indent_pt / 72.0) if indent_pt else Inches(0)
        for r in par.runs:
            r.font.name = body_font; r.font.size = Pt(body_size)

    if title:
        h = doc.add_heading(_sanitize_text(title, ensure_font=False), level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    fig_no = 1; tbl_no = 1
    for kind, text in _parse_md(md):
        if kind == "blank":
            continue
        elif kind == "h1":
            doc.add_heading(_sanitize_text(_strip_inline(text), ensure_font=False), level=1)
        elif kind == "h2":
            doc.add_heading(_sanitize_text(_strip_inline(text), ensure_font=False), level=2)
        elif kind == "h3":
            doc.add_heading(_sanitize_text(_strip_inline(text), ensure_font=False), level=3)
        elif kind == "li":
            par = doc.add_paragraph(_sanitize_text(_strip_inline(text), ensure_font=False), style="List Bullet")
            par.paragraph_format.line_spacing = line_spacing
        elif kind == "table":
            rows = text
            if rows and len(rows) >= 1:
                ncol = max(len(r) for r in rows)
                t = doc.add_table(rows=0, cols=ncol)
                try:
                    t.style = "Light Grid Accent 1"
                except Exception:
                    try:
                        t.style = "Table Grid"
                    except Exception:
                        pass
                for ri, row in enumerate(rows):
                    cells = t.add_row().cells
                    for ci in range(ncol):
                        val = row[ci] if ci < len(row) else ""
                        cells[ci].text = _sanitize_text(_strip_inline(val), ensure_font=False)
                        if ri == 0:
                            for pp in cells[ci].paragraphs:
                                for rr in pp.runs:
                                    rr.font.bold = True
        elif kind == "img":
            url, cap = text
            local = (url if os.path.exists(url)
                     else (figures.get(url) if figures and os.path.exists(str(figures.get(url, ""))) else None))
            if local and os.path.exists(local):
                try:
                    doc.add_picture(local, width=Inches(5.5))
                    if doc.paragraphs:
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if cap:
                        c = doc.add_paragraph(_sanitize_text(cap, ensure_font=False))
                        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if c.runs:
                            c.runs[0].italic = True; c.runs[0].font.size = Pt(max(8, body_size - 1))
                except Exception:
                    doc.add_paragraph(f"[Figure: {cap or url}]")
        else:
            par = doc.add_paragraph(_sanitize_text(_strip_inline(text), ensure_font=False))
            _style_body(par)
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    try:
        _m = doc_meta or {}
        cp = doc.core_properties
        cp.title = _sanitize_text(_m.get("title") or title or "", ensure_font=False)
        cp.author = _sanitize_text(_m.get("author") or "ILMA", ensure_font=False)
        cp.subject = _sanitize_text(_m.get("subject") or "", ensure_font=False)
        cp.keywords = _sanitize_text(_m.get("keywords") or "", ensure_font=False)
    except Exception:
        pass
    doc.save(out_path)
    return out_path


def to_pdf(md: str, out_path: str, title: str = None,
           figures: Optional[Dict[str, str]] = None, doc_type: str = "report",
           doc_meta: Optional[Dict[str, str]] = None) -> str:
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib.enums import TA_JUSTIFY, TA_LEFT, TA_CENTER, TA_RIGHT
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Image as RLImage, ListFlowable, ListItem, Table, TableStyle, KeepTogether)
    from reportlab.lib import colors
    Path(out_path).parent.mkdir(parents=True, exist_ok=True)
    _font = _register_unicode_font()
    base_font = _font or "Helvetica"
    bold_font = _PDF_FONT_BOLD or "Helvetica-Bold"
    ital_font = _PDF_FONT_ITALIC or "Helvetica-Oblique"

    fp = _fmt_profile(doc_type)
    _ALN = {"justify": TA_JUSTIFY, "left": TA_LEFT, "center": TA_CENTER, "right": TA_RIGHT}
    align = _ALN.get(fp.get("align", "justify"), TA_JUSTIFY)
    body_size = float(fp.get("body_size", 11) or 11)
    leading = body_size * float(fp.get("line_spacing", 1.15) or 1.15)
    space_after = float(fp.get("para_space_after", 8) or 0)
    indent = float(fp.get("first_line_indent", 0) or 0)
    block = bool(fp.get("block_paragraphs"))
    margin = float(fp.get("margin_in", 1.0) or 1.0) * inch

    styles = getSampleStyleSheet()
    for _sn in ("Title", "Heading1", "Heading2", "Heading3", "BodyText", "Normal"):
        try:
            styles[_sn].fontName = (bold_font if (_sn == "Title" or _sn.startswith("Heading")) else base_font)
        except Exception:
            pass
    body = ParagraphStyle("BodyFmt", parent=styles["BodyText"], fontName=base_font,
                          fontSize=body_size, leading=leading, alignment=align,
                          spaceAfter=(space_after if block else 2),
                          firstLineIndent=(0 if block else indent))
    cap_style = ParagraphStyle("Caption", parent=body, fontName=ital_font,
                               fontSize=max(8, body_size - 1), alignment=TA_CENTER, spaceAfter=8)
    story = []
    if title:
        story.append(Paragraph(html.escape(_sanitize_text(title)), styles["Title"]))
        story.append(Spacer(1, 0.2 * inch))
    bullets = []

    def flush_bullets():
        nonlocal bullets
        if bullets:
            story.append(ListFlowable([ListItem(Paragraph(b, body)) for b in bullets], bulletType="bullet"))
            bullets = []

    for kind, text in _parse_md(md):
        if kind == "li":
            bullets.append(html.escape(_sanitize_text(_strip_inline(text))))
            continue
        flush_bullets()
        if kind == "blank":
            story.append(Spacer(1, 4))
        elif kind == "h1":
            story.append(Paragraph(html.escape(_sanitize_text(_strip_inline(text))), styles["Heading1"]))
        elif kind == "h2":
            story.append(Paragraph(html.escape(_sanitize_text(_strip_inline(text))), styles["Heading2"]))
        elif kind == "h3":
            story.append(Paragraph(html.escape(_sanitize_text(_strip_inline(text))), styles["Heading3"]))
        elif kind == "table":
            rows = text
            if rows:
                ncol = max(len(r) for r in rows)
                cell_style = ParagraphStyle("Cell", parent=body, fontSize=max(8, body_size - 1),
                                            alignment=TA_LEFT, firstLineIndent=0, spaceAfter=0,
                                            leading=(body_size) * 1.1)
                head_style = ParagraphStyle("CellHead", parent=cell_style, fontName=bold_font)
                data = []
                for ri, row in enumerate(rows):
                    cells = []
                    for ci in range(ncol):
                        val = html.escape(_sanitize_text(_strip_inline(row[ci] if ci < len(row) else "")))
                        cells.append(Paragraph(val, head_style if ri == 0 else cell_style))
                    data.append(cells)
                from reportlab.lib.pagesizes import A4
                avail = A4[0] - 2 * margin
                t = Table(data, colWidths=[avail / ncol] * ncol, hAlign="CENTER")
                t.setStyle(TableStyle([
                    ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f4e79")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f2f2f2")]),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5), ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 3), ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]))
                story.append(Spacer(1, 4)); story.append(t); story.append(Spacer(1, 8))
        elif kind == "img":
            url, cap = text
            local = (url if os.path.exists(url)
                     else (figures.get(url) if figures and os.path.exists(str(figures.get(url, ""))) else None))
            if local and os.path.exists(local):
                try:
                    _imgflow = [RLImage(local, width=5.0 * inch, height=3.2 * inch, kind="proportional")]
                    if cap:
                        _imgflow.append(Paragraph(html.escape(_sanitize_text(cap)), cap_style))
                    story.append(KeepTogether(_imgflow))
                except Exception:
                    story.append(Paragraph(f"[Figure: {html.escape(_sanitize_text(cap or url))}]", body))
            else:
                story.append(Paragraph(f"[Figure: {html.escape(_sanitize_text(cap or url))}]", body))
        else:
            story.append(Paragraph(html.escape(_sanitize_text(_strip_inline(text))), body))
    flush_bullets()

    _meta = doc_meta or {}
    _pg = _page_size(fp.get("page_size", "A4"))

    def _on_page(canvas, doc_):
        canvas.saveState()
        # PDF document metadata (first page is enough; reportlab applies doc-wide)
        try:
            canvas.setTitle(_sanitize_text(_meta.get("title") or title or ""))
            if _meta.get("author"):
                canvas.setAuthor(_sanitize_text(_meta.get("author")))
            if _meta.get("subject"):
                canvas.setSubject(_sanitize_text(_meta.get("subject")))
            if _meta.get("keywords"):
                canvas.setKeywords(_sanitize_text(_meta.get("keywords")))
        except Exception:
            pass
        # footer page number, centered
        try:
            canvas.setFont(base_font, 9)
            canvas.setFillColorRGB(0.4, 0.4, 0.4)
            canvas.drawCentredString(_pg[0] / 2.0, margin * 0.45, str(doc_.page))
        except Exception:
            pass
        canvas.restoreState()

    SimpleDocTemplate(out_path, pagesize=_pg,
                      leftMargin=margin, rightMargin=margin,
                      topMargin=margin, bottomMargin=margin,
                      title=_sanitize_text(_meta.get("title") or title or ""),
                      author=_sanitize_text(_meta.get("author") or "ILMA"),
                      subject=_sanitize_text(_meta.get("subject") or ""),
                      keywords=_sanitize_text(_meta.get("keywords") or "")
                      ).build(story, onFirstPage=_on_page, onLaterPages=_on_page)
    return out_path


def export(markdown_text: str, out_basepath: str,
           formats=("docx", "pdf", "html"), title: str = None,
           figures: Optional[Dict[str, str]] = None, doc_type: str = "report",
           doc_meta: Optional[Dict[str, str]] = None) -> Dict[str, str]:
    base = Path(out_basepath)
    base.parent.mkdir(parents=True, exist_ok=True)
    # normalize formats: accept "docx,pdf" string or list/tuple
    if isinstance(formats, str):
        formats = [f.strip() for f in formats.replace(";", ",").split(",") if f.strip()]
    out = {}
    # always save canonical markdown
    md_path = str(base.with_suffix(".md"))
    Path(md_path).write_text(markdown_text)
    out["md"] = md_path
    for fmt in formats:
        try:
            if fmt == "html":
                p = str(base.with_suffix(".html"))
                Path(p).write_text(to_html(markdown_text, title or base.stem, doc_type=doc_type))
                out["html"] = p
            elif fmt == "docx":
                out["docx"] = to_docx(markdown_text, str(base.with_suffix(".docx")), title, figures, doc_type=doc_type, doc_meta=doc_meta)
            elif fmt == "pdf":
                out["pdf"] = to_pdf(markdown_text, str(base.with_suffix(".pdf")), title, figures, doc_type=doc_type, doc_meta=doc_meta)
        except Exception as e:
            out[fmt + "_error"] = str(e)[:160]
    return out


if __name__ == "__main__":
    import sys, json
    sample = "# Test Doc\n\n## Intro\n\nThis is a **bold** test with a fact [1].\n\n- point one\n- point two\n\n## References\n\n[1] Example source."
    r = export(sample, "/tmp/ilma_export_test", title="Test Doc")
    print(json.dumps(r, indent=2))
